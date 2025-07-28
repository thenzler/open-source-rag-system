"""
Document Processing Service
Handles extraction, chunking, and preprocessing of various document formats.
"""

import asyncio
import hashlib
import logging
import mimetypes
import os
import shutil
import tempfile
import uuid
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any

import aiofiles
import pandas as pd
import pymupdf  # PyMuPDF (fitz)
import pdfplumber
from bs4 import BeautifulSoup
from celery import Celery
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from lxml import etree
import openpyxl
from PyPDF2 import PdfReader
import pytesseract
from PIL import Image
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import get_settings
from app.core.database import get_database
from app.models.documents import Document, DocumentChunk, ProcessingStatus
from app.services.vector_service import VectorService
from app.core.exceptions import ProcessingError, UnsupportedFormatError

logger = logging.getLogger(__name__)
settings = get_settings()

# Initialize Celery for background processing
celery_app = Celery(
    'document_processor',
    broker=settings.celery_broker,
    backend=settings.celery_result_backend
)

# Initialize vector service
vector_service = VectorService()


class DocumentProcessor:
    """Main document processing service."""
    
    def __init__(self):
        self.supported_formats = {
            'application/pdf': self._process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._process_xlsx,
            'application/xml': self._process_xml,
            'text/xml': self._process_xml,
            'text/plain': self._process_text,
            'text/markdown': self._process_text,
            'text/csv': self._process_csv,
        }
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    async def process_document(
        self, 
        file_path: str, 
        document_id: str, 
        mime_type: str,
        db: AsyncSession
    ) -> Dict[str, Any]:
        """
        Process a document: extract text, chunk it, and store vectors.
        """
        try:
            logger.info(f"Starting processing for document {document_id}")
            
            # Update status to processing
            await self._update_processing_status(
                document_id, 
                ProcessingStatus.PROCESSING, 
                db,
                progress=0
            )
            
            # Extract text and metadata
            extraction_result = await self._extract_content(file_path, mime_type)
            
            await self._update_processing_status(
                document_id, 
                ProcessingStatus.PROCESSING, 
                db,
                progress=30,
                message="Text extraction completed"
            )
            
            # Chunk the text
            chunks = await self._chunk_text(
                extraction_result['text'],
                extraction_result.get('metadata', {})
            )
            
            await self._update_processing_status(
                document_id, 
                ProcessingStatus.PROCESSING, 
                db,
                progress=50,
                message="Text chunking completed"
            )
            
            # Store chunks in database
            chunk_records = await self._store_chunks(document_id, chunks, db)
            
            await self._update_processing_status(
                document_id, 
                ProcessingStatus.PROCESSING, 
                db,
                progress=70,
                message="Chunks stored in database"
            )
            
            # Generate embeddings and store in vector database
            await self._generate_and_store_embeddings(document_id, chunk_records)
            
            await self._update_processing_status(
                document_id, 
                ProcessingStatus.PROCESSING, 
                db,
                progress=90,
                message="Embeddings generated and stored"
            )
            
            # Update final status
            await self._update_processing_status(
                document_id, 
                ProcessingStatus.COMPLETED, 
                db,
                progress=100,
                message="Processing completed successfully"
            )
            
            result = {
                'status': 'completed',
                'total_chunks': len(chunks),
                'total_characters': len(extraction_result['text']),
                'metadata': extraction_result.get('metadata', {}),
                'processing_time_ms': extraction_result.get('processing_time_ms', 0)
            }
            
            logger.info(f"Document {document_id} processed successfully: {result}")
            return result
            
        except Exception as e:
            logger.error(f"Error processing document {document_id}: {e}")
            await self._update_processing_status(
                document_id, 
                ProcessingStatus.FAILED, 
                db,
                message=f"Processing failed: {str(e)}"
            )
            raise ProcessingError(f"Document processing failed: {str(e)}")
    
    async def _extract_content(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """Extract text content from a file based on its MIME type."""
        if mime_type not in self.supported_formats:
            raise UnsupportedFormatError(f"Unsupported file format: {mime_type}")
        
        processor = self.supported_formats[mime_type]
        return await processor(file_path)
    
    async def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF files with fallback methods."""
        text = ""
        metadata = {}
        
        try:
            # Method 1: Try PyMuPDF (fastest and most accurate)
            doc = pymupdf.open(file_path)
            pages_text = []
            
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                if page_text.strip():
                    pages_text.append({
                        'page_number': page_num + 1,
                        'text': page_text,
                        'char_count': len(page_text)
                    })
            
            text = '\n'.join([p['text'] for p in pages_text])
            
            # Extract metadata
            metadata.update({
                'pages': len(doc),
                'page_details': pages_text[:5],  # Store first 5 pages details
                'title': doc.metadata.get('title', ''),
                'author': doc.metadata.get('author', ''),
                'subject': doc.metadata.get('subject', ''),
                'creator': doc.metadata.get('creator', ''),
                'producer': doc.metadata.get('producer', ''),
                'creation_date': str(doc.metadata.get('creationDate', '')),
                'modification_date': str(doc.metadata.get('modDate', '')),
                'extraction_method': 'pymupdf'
            })
            
            doc.close()
            
            # If no text found, try OCR
            if not text.strip() and settings.enable_ocr:
                logger.info(f"No text found in PDF, attempting OCR for {file_path}")
                text, ocr_metadata = await self._ocr_pdf(file_path)
                metadata.update(ocr_metadata)
                metadata['extraction_method'] = 'ocr'
            
        except Exception as e:
            logger.warning(f"PyMuPDF failed for {file_path}: {e}, trying pdfplumber")
            
            try:
                # Method 2: Fallback to pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    pages_text = []
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            pages_text.append({
                                'page_number': page_num + 1,
                                'text': page_text,
                                'char_count': len(page_text)
                            })
                    
                    text = '\n'.join([p['text'] for p in pages_text])
                    metadata.update({
                        'pages': len(pdf.pages),
                        'page_details': pages_text[:5],
                        'extraction_method': 'pdfplumber'
                    })
                    
            except Exception as e2:
                logger.warning(f"pdfplumber failed for {file_path}: {e2}, trying PyPDF2")
                
                try:
                    # Method 3: Final fallback to PyPDF2
                    with open(file_path, 'rb') as file:
                        pdf_reader = PdfReader(file)
                        pages_text = []
                        
                        for page_num, page in enumerate(pdf_reader.pages):
                            page_text = page.extract_text()
                            if page_text.strip():
                                pages_text.append({
                                    'page_number': page_num + 1,
                                    'text': page_text,
                                    'char_count': len(page_text)
                                })
                        
                        text = '\n'.join([p['text'] for p in pages_text])
                        
                        # Extract metadata
                        pdf_metadata = pdf_reader.metadata or {}
                        metadata.update({
                            'pages': len(pdf_reader.pages),
                            'page_details': pages_text[:5],
                            'title': pdf_metadata.get('/Title', ''),
                            'author': pdf_metadata.get('/Author', ''),
                            'subject': pdf_metadata.get('/Subject', ''),
                            'creator': pdf_metadata.get('/Creator', ''),
                            'producer': pdf_metadata.get('/Producer', ''),
                            'creation_date': str(pdf_metadata.get('/CreationDate', '')),
                            'modification_date': str(pdf_metadata.get('/ModDate', '')),
                            'extraction_method': 'pypdf2'
                        })
                        
                except Exception as e3:
                    raise ProcessingError(f"All PDF extraction methods failed: {e3}")
        
        return {
            'text': text,
            'metadata': metadata,
            'format': 'pdf'
        }
    
    async def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from Word documents."""
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):  # Only add non-empty rows
                        table_data.append(row_data)
                if table_data:
                    tables_text.append('\n'.join(['\t'.join(row) for row in table_data]))
            
            # Combine all text
            all_text = '\n'.join(paragraphs)
            if tables_text:
                all_text += '\n\n' + '\n\n'.join(tables_text)
            
            # Extract metadata
            metadata = {
                'paragraphs_count': len(paragraphs),
                'tables_count': len(doc.tables),
                'sections_count': len(doc.sections),
                'extraction_method': 'python-docx',
                'format': 'docx'
            }
            
            # Core properties
            if hasattr(doc, 'core_properties'):
                core_props = doc.core_properties
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'keywords': core_props.keywords or '',
                    'comments': core_props.comments or '',
                    'created': str(core_props.created) if core_props.created else '',
                    'modified': str(core_props.modified) if core_props.modified else '',
                    'last_modified_by': core_props.last_modified_by or ''
                })
            
            return {
                'text': all_text,
                'metadata': metadata,
                'format': 'docx'
            }
            
        except Exception as e:
            raise ProcessingError(f"Failed to process DOCX file: {e}")
    
    async def _process_xlsx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from Excel files."""
        try:
            # Load workbook
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            sheets_data = []
            total_rows = 0
            total_cells = 0
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Convert sheet to DataFrame for easier processing
                data = []
                for row in sheet.iter_rows(values_only=True):
                    # Filter out None values and convert to strings
                    row_data = [str(cell) if cell is not None else '' for cell in row]
                    if any(row_data):  # Only include non-empty rows
                        data.append(row_data)
                        total_cells += len([cell for cell in row_data if cell])
                
                total_rows += len(data)
                
                if data:
                    # Convert to text representation
                    sheet_text = f"Sheet: {sheet_name}\n"
                    sheet_text += '\n'.join(['\t'.join(row) for row in data])
                    sheets_data.append({
                        'sheet_name': sheet_name,
                        'text': sheet_text,
                        'rows': len(data),
                        'columns': len(data[0]) if data else 0
                    })
            
            # Combine all sheets
            all_text = '\n\n'.join([sheet['text'] for sheet in sheets_data])
            
            # Metadata
            metadata = {
                'sheets_count': len(workbook.sheetnames),
                'sheet_names': workbook.sheetnames,
                'total_rows': total_rows,
                'total_cells': total_cells,
                'sheets_data': sheets_data,
                'extraction_method': 'openpyxl',
                'format': 'xlsx'
            }
            
            workbook.close()
            
            return {
                'text': all_text,
                'metadata': metadata,
                'format': 'xlsx'
            }
            
        except Exception as e:
            raise ProcessingError(f"Failed to process XLSX file: {e}")
    
    async def _process_xml(self, file_path: str) -> Dict[str, Any]:
        """Extract text from XML files."""
        try:
            # Parse XML
            tree = etree.parse(file_path)
            root = tree.getroot()
            
            # Extract all text content
            text_content = []
            
            def extract_text(element):
                if element.text and element.text.strip():
                    text_content.append(element.text.strip())
                for child in element:
                    extract_text(child)
                if element.tail and element.tail.strip():
                    text_content.append(element.tail.strip())
            
            extract_text(root)
            
            # Also try BeautifulSoup for HTML-like XML
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                soup = BeautifulSoup(content, 'xml')
                soup_text = soup.get_text(separator=' ', strip=True)
                if len(soup_text) > len(' '.join(text_content)):
                    text_content = [soup_text]
            except:
                pass  # Fallback to lxml extraction
            
            all_text = ' '.join(text_content)
            
            # Metadata
            metadata = {
                'root_tag': root.tag,
                'namespace': root.nsmap if hasattr(root, 'nsmap') else {},
                'elements_count': len(list(root.iter())),
                'extraction_method': 'lxml',
                'format': 'xml'
            }
            
            return {
                'text': all_text,
                'metadata': metadata,
                'format': 'xml'
            }
            
        except Exception as e:
            raise ProcessingError(f"Failed to process XML file: {e}")
    
    async def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from plain text files."""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = await f.read()
            
            # Basic text statistics
            lines = text.split('\n')
            words = text.split()
            
            metadata = {
                'lines_count': len(lines),
                'words_count': len(words),
                'characters_count': len(text),
                'encoding': 'utf-8',
                'extraction_method': 'direct',
                'format': 'text'
            }
            
            return {
                'text': text,
                'metadata': metadata,
                'format': 'text'
            }
            
        except Exception as e:
            raise ProcessingError(f"Failed to process text file: {e}")
    
    async def _process_csv(self, file_path: str) -> Dict[str, Any]:
        """Extract text from CSV files."""
        try:
            # Read CSV
            df = pd.read_csv(file_path, encoding='utf-8', errors='ignore')
            
            # Convert to text representation
            text_parts = []
            
            # Add header
            text_parts.append("Columns: " + ", ".join(df.columns.tolist()))
            
            # Add data rows (limit for large files)
            max_rows = 1000  # Prevent memory issues
            for idx, row in df.head(max_rows).iterrows():
                row_text = []
                for col, value in row.items():
                    if pd.notna(value):
                        row_text.append(f"{col}: {value}")
                if row_text:
                    text_parts.append("; ".join(row_text))
            
            all_text = '\n'.join(text_parts)
            
            # Metadata
            metadata = {
                'rows_count': len(df),
                'columns_count': len(df.columns),
                'column_names': df.columns.tolist(),
                'data_types': df.dtypes.to_dict(),
                'null_counts': df.isnull().sum().to_dict(),
                'extraction_method': 'pandas',
                'format': 'csv',
                'sample_rows': min(max_rows, len(df))
            }
            
            return {
                'text': all_text,
                'metadata': metadata,
                'format': 'csv'
            }
            
        except Exception as e:
            raise ProcessingError(f"Failed to process CSV file: {e}")
    
    async def _ocr_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """Perform OCR on PDF file."""
        try:
            doc = pymupdf.open(file_path)
            ocr_text = []
            
            for page_num, page in enumerate(doc):
                # Convert page to image
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")
                
                # Save to temporary file
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_img:
                    temp_img.write(img_data)
                    temp_img_path = temp_img.name
                
                try:
                    # Perform OCR
                    image = Image.open(temp_img_path)
                    page_text = pytesseract.image_to_string(
                        image, 
                        lang=settings.ocr_language,
                        config=settings.tesseract_config
                    )
                    
                    if page_text.strip():
                        ocr_text.append(f"Page {page_num + 1}:\n{page_text}")
                
                finally:
                    # Clean up temp file
                    os.unlink(temp_img_path)
            
            doc.close()
            
            text = '\n\n'.join(ocr_text)
            metadata = {
                'ocr_pages': len(ocr_text),
                'ocr_language': settings.ocr_language,
                'extraction_method': 'ocr'
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"OCR failed: {e}")
            return "", {"ocr_error": str(e)}
    
    async def _chunk_text(self, text: str, metadata: Dict) -> List[Dict[str, Any]]:
        """Split text into chunks for processing."""
        if not text.strip():
            return []
        
        # Use RecursiveCharacterTextSplitter for intelligent chunking
        chunks = self.text_splitter.split_text(text)
        
        # Create chunk objects with metadata
        chunk_objects = []
        for i, chunk_text in enumerate(chunks):
            if len(chunk_text.strip()) >= settings.min_chunk_size:
                chunk_objects.append({
                    'index': i,
                    'text': chunk_text,
                    'char_count': len(chunk_text),
                    'word_count': len(chunk_text.split()),
                    'metadata': {
                        'source_format': metadata.get('format', 'unknown'),
                        'chunk_index': i,
                        'total_chunks': len(chunks),
                        'extraction_method': metadata.get('extraction_method', 'unknown')
                    }
                })
        
        return chunk_objects
    
    async def _store_chunks(
        self, 
        document_id: str, 
        chunks: List[Dict[str, Any]], 
        db: AsyncSession
    ) -> List[DocumentChunk]:
        """Store text chunks in the database."""
        chunk_records = []
        
        for chunk_data in chunks:
            chunk = DocumentChunk(
                id=str(uuid.uuid4()),
                document_id=document_id,
                chunk_index=chunk_data['index'],
                content=chunk_data['text'],
                chunk_metadata=chunk_data['metadata'],
                char_count=chunk_data['char_count'],
                word_count=chunk_data['word_count']
            )
            
            db.add(chunk)
            chunk_records.append(chunk)
        
        await db.commit()
        return chunk_records
    
    async def _generate_and_store_embeddings(
        self, 
        document_id: str, 
        chunks: List[DocumentChunk]
    ):
        """Generate embeddings for chunks and store in vector database."""
        try:
            # Prepare data for vector storage
            texts = [chunk.content for chunk in chunks]
            metadatas = []
            
            for chunk in chunks:
                metadata = {
                    'document_id': document_id,
                    'chunk_id': chunk.id,
                    'chunk_index': chunk.chunk_index,
                    'char_count': chunk.char_count,
                    'word_count': chunk.word_count
                }
                metadata.update(chunk.chunk_metadata or {})
                metadatas.append(metadata)
            
            # Store in vector database
            vector_ids = await vector_service.add_documents(
                texts=texts,
                metadatas=metadatas,
                document_id=document_id
            )
            
            # Update chunks with vector IDs
            for chunk, vector_id in zip(chunks, vector_ids):
                chunk.vector_id = vector_id
            
            logger.info(f"Generated and stored {len(vector_ids)} embeddings for document {document_id}")
            
        except Exception as e:
            logger.error(f"Failed to generate embeddings for document {document_id}: {e}")
            raise ProcessingError(f"Embedding generation failed: {e}")
    
    async def _update_processing_status(
        self,
        document_id: str,
        status: ProcessingStatus,
        db: AsyncSession,
        progress: Optional[int] = None,
        message: Optional[str] = None
    ):
        """Update document processing status."""
        try:
            # Find document
            result = await db.execute(
                "SELECT * FROM documents WHERE id = :document_id",
                {"document_id": document_id}
            )
            document = result.fetchone()
            
            if document:
                update_data = {"status": status.value}
                if progress is not None:
                    update_data["progress"] = progress
                if message:
                    update_data["status_message"] = message
                
                await db.execute(
                    "UPDATE documents SET status = :status, progress = :progress, status_message = :message WHERE id = :document_id",
                    {
                        "status": status.value,
                        "progress": progress,
                        "message": message,
                        "document_id": document_id
                    }
                )
                await db.commit()
                
        except Exception as e:
            logger.error(f"Failed to update status for document {document_id}: {e}")


# Celery task for background processing
@celery_app.task(bind=True)
def process_document_task(self, document_id: str, file_path: str, mime_type: str):
    """Celery task for processing documents in the background."""
    try:
        processor = DocumentProcessor()
        
        # Run async processing in event loop
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        
        async def process():
            async for db in get_database():
                return await processor.process_document(file_path, document_id, mime_type, db)
        
        result = loop.run_until_complete(process())
        loop.close()
        
        return result
        
    except Exception as e:
        logger.error(f"Celery task failed for document {document_id}: {e}")
        # Update status to failed
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        
        async def update_failed_status():
            async for db in get_database():
                processor = DocumentProcessor()
                await processor._update_processing_status(
                    document_id, 
                    ProcessingStatus.FAILED, 
                    db,
                    message=f"Processing failed: {str(e)}"
                )
        
        loop.run_until_complete(update_failed_status())
        loop.close()
        
        raise


# Service class for API integration
class DocumentService:
    """Service class for document operations."""
    
    def __init__(self):
        self.processor = DocumentProcessor()
    
    async def initialize(self):
        """Initialize the document service."""
        # Ensure upload directories exist
        os.makedirs(settings.upload_directory, exist_ok=True)
        os.makedirs(settings.processed_directory, exist_ok=True)
        
        # Initialize vector service
        await vector_service.initialize()
    
    async def upload_document(
        self, 
        file, 
        metadata: Optional[str], 
        user_id: str, 
        db: AsyncSession
    ) -> Document:
        """Handle document upload and initiate processing."""
        try:
            # Generate document ID
            document_id = str(uuid.uuid4())
            
            # Calculate file hash
            file_content = await file.read()
            await file.seek(0)  # Reset file pointer
            file_hash = hashlib.sha256(file_content).hexdigest()
            
            # Save file
            file_extension = Path(file.filename).suffix
            stored_filename = f"{document_id}{file_extension}"
            file_path = os.path.join(settings.upload_directory, stored_filename)
            
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(file_content)
            
            # Detect MIME type
            mime_type, _ = mimetypes.guess_type(file.filename)
            if not mime_type:
                mime_type = file.content_type
            
            # Create database record
            document = Document(
                id=document_id,
                filename=file.filename,
                file_path=file_path,
                mime_type=mime_type,
                file_size=len(file_content),
                checksum=file_hash,
                user_id=user_id,
                status=ProcessingStatus.PENDING,
                metadata=metadata
            )
            
            db.add(document)
            await db.commit()
            await db.refresh(document)
            
            return document
            
        except Exception as e:
            logger.error(f"Document upload failed: {e}")
            raise ProcessingError(f"Upload failed: {e}")
    
    async def process_document_async(self, document_id: str, db: AsyncSession):
        """Start background processing of a document."""
        try:
            # Get document info
            result = await db.execute(
                "SELECT file_path, mime_type FROM documents WHERE id = :document_id",
                {"document_id": document_id}
            )
            document_info = result.fetchone()
            
            if not document_info:
                raise DocumentNotFoundError(f"Document {document_id} not found")
            
            # Queue Celery task
            process_document_task.delay(
                document_id=document_id,
                file_path=document_info.file_path,
                mime_type=document_info.mime_type
            )
            
        except Exception as e:
            logger.error(f"Failed to queue processing for document {document_id}: {e}")
            raise
    
    async def health_check(self) -> bool:
        """Check if document service is healthy."""
        try:
            # Check if upload directory is writable
            test_file = os.path.join(settings.upload_directory, ".health_check")
            with open(test_file, 'w') as f:
                f.write("test")
            os.remove(test_file)
            
            return True
        except Exception:
            return False
