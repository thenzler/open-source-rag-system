"""
Document Processor Service - Handles document parsing, text extraction, and chunking
"""

import asyncio
import hashlib
import logging
import os
from pathlib import Path
from typing import Dict, List, Optional, Any
import uuid
from datetime import datetime

import aiofiles
import aiofiles.os
from fastapi import FastAPI, UploadFile, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from celery import Celery
from pydantic import BaseModel
import uvicorn

# Document processing libraries
import PyPDF2
import docx
import openpyxl
from bs4 import BeautifulSoup
import xml.etree.ElementTree as ET
import pandas as pd
from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangChainDocument

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Environment variables
UPLOAD_DIRECTORY = os.getenv("UPLOAD_DIRECTORY", "/app/storage/uploads")
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "512"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "50"))
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "sentence-transformers/all-mpnet-base-v2")
CELERY_BROKER = os.getenv("CELERY_BROKER", "redis://localhost:6379/0")
DATABASE_URL = os.getenv("DATABASE_URL", "postgresql://raguser:password@localhost:5432/ragdb")
QDRANT_URL = os.getenv("QDRANT_URL", "http://localhost:6333")

# Initialize Celery
celery_app = Celery(
    "document_processor",
    broker=CELERY_BROKER,
    backend=CELERY_BROKER
)

# Pydantic models
class DocumentProcessingRequest(BaseModel):
    document_id: str
    file_path: str
    file_type: str
    metadata: Optional[Dict[str, Any]] = None

class DocumentChunk(BaseModel):
    chunk_id: str
    document_id: str
    content: str
    metadata: Dict[str, Any]
    chunk_index: int
    embedding: Optional[List[float]] = None

class ProcessingResult(BaseModel):
    document_id: str
    status: str
    chunks_count: int
    processing_time: float
    error: Optional[str] = None

# FastAPI app
app = FastAPI(
    title="Document Processor Service",
    description="Handles document parsing, text extraction, and chunking",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Global variables for services
embedding_model = None
text_splitter = None

@app.on_event("startup")
async def startup_event():
    """Initialize services on startup"""
    global embedding_model, text_splitter
    
    logger.info("Initializing Document Processor Service...")
    
    # Create upload directory
    Path(UPLOAD_DIRECTORY).mkdir(parents=True, exist_ok=True)
    
    # Initialize embedding model
    try:
        embedding_model = SentenceTransformer(EMBEDDING_MODEL)
        logger.info(f"Loaded embedding model: {EMBEDDING_MODEL}")
    except Exception as e:
        logger.error(f"Failed to load embedding model: {e}")
        embedding_model = None
    
    # Initialize text splitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    logger.info("Document Processor Service initialized successfully")

class DocumentProcessor:
    """Document processing utilities"""
    
    SUPPORTED_FORMATS = {
        'pdf': ['pdf'],
        'word': ['docx', 'doc'],
        'excel': ['xlsx', 'xls'],
        'text': ['txt', 'md', 'rtf'],
        'html': ['html', 'htm'],
        'xml': ['xml'],
        'csv': ['csv']
    }
    
    @staticmethod
    def get_file_type(filename: str) -> str:
        """Determine file type from filename"""
        extension = filename.lower().split('.')[-1]
        
        for file_type, extensions in DocumentProcessor.SUPPORTED_FORMATS.items():
            if extension in extensions:
                return file_type
        
        return 'unknown'
    
    @staticmethod
    def validate_file(file: UploadFile) -> bool:
        """Validate uploaded file"""
        if not file.filename:
            return False
            
        file_type = DocumentProcessor.get_file_type(file.filename)
        return file_type != 'unknown'
    
    @staticmethod
    async def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise
    
    @staticmethod
    async def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise
    
    @staticmethod
    async def extract_text_from_excel(file_path: str) -> str:
        """Extract text from Excel file"""
        try:
            workbook = openpyxl.load_workbook(file_path)
            text = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"Sheet: {sheet_name}\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                    text += row_text + "\n"
                text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from Excel: {e}")
            raise
    
    @staticmethod
    async def extract_text_from_html(file_path: str) -> str:
        """Extract text from HTML file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                text = soup.get_text()
                # Clean up text
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
                return text
        except Exception as e:
            logger.error(f"Error extracting text from HTML: {e}")
            raise
    
    @staticmethod
    async def extract_text_from_xml(file_path: str) -> str:
        """Extract text from XML file"""
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            def extract_text_recursive(element):
                text = element.text or ""
                for child in element:
                    text += extract_text_recursive(child)
                text += element.tail or ""
                return text
            
            return extract_text_recursive(root).strip()
        except Exception as e:
            logger.error(f"Error extracting text from XML: {e}")
            raise
    
    @staticmethod
    async def extract_text_from_csv(file_path: str) -> str:
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(file_path)
            text = ""
            
            # Add column headers
            text += "Columns: " + ", ".join(df.columns) + "\n\n"
            
            # Convert dataframe to text
            for index, row in df.iterrows():
                row_text = "\t".join([str(value) for value in row.values])
                text += row_text + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from CSV: {e}")
            raise
    
    @staticmethod
    async def extract_text_from_text(file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                return await file.read()
        except Exception as e:
            logger.error(f"Error extracting text from text file: {e}")
            raise
    
    @staticmethod
    async def extract_text(file_path: str, file_type: str) -> str:
        """Extract text from file based on type"""
        extractors = {
            'pdf': DocumentProcessor.extract_text_from_pdf,
            'word': DocumentProcessor.extract_text_from_docx,
            'excel': DocumentProcessor.extract_text_from_excel,
            'html': DocumentProcessor.extract_text_from_html,
            'xml': DocumentProcessor.extract_text_from_xml,
            'csv': DocumentProcessor.extract_text_from_csv,
            'text': DocumentProcessor.extract_text_from_text
        }
        
        extractor = extractors.get(file_type)
        if not extractor:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        return await extractor(file_path)
    
    @staticmethod
    def create_chunks(text: str, document_id: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Create chunks from text"""
        if not text_splitter:
            raise RuntimeError("Text splitter not initialized")
        
        # Create langchain document
        doc = LangChainDocument(page_content=text, metadata=metadata)
        
        # Split into chunks
        chunks = text_splitter.split_documents([doc])
        
        # Convert to our format
        document_chunks = []
        for i, chunk in enumerate(chunks):
            chunk_id = str(uuid.uuid4())
            chunk_metadata = {
                **metadata,
                'chunk_index': i,
                'chunk_id': chunk_id,
                'document_id': document_id
            }
            
            document_chunks.append(DocumentChunk(
                chunk_id=chunk_id,
                document_id=document_id,
                content=chunk.page_content,
                metadata=chunk_metadata,
                chunk_index=i
            ))
        
        return document_chunks
    
    @staticmethod
    def generate_embeddings(chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """Generate embeddings for chunks"""
        if not embedding_model:
            logger.warning("Embedding model not available, skipping embeddings")
            return chunks
        
        try:
            texts = [chunk.content for chunk in chunks]
            embeddings = embedding_model.encode(texts)
            
            for chunk, embedding in zip(chunks, embeddings):
                chunk.embedding = embedding.tolist()
            
            return chunks
        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")
            return chunks

# Celery tasks
@celery_app.task(bind=True)
def process_document_task(self, document_id: str, file_path: str, file_type: str, metadata: Dict[str, Any] = None):
    """Process document in background"""
    try:
        logger.info(f"Processing document {document_id} of type {file_type}")
        
        # Extract text
        text = asyncio.run(DocumentProcessor.extract_text(file_path, file_type))
        
        if not text.strip():
            raise ValueError("No text extracted from document")
        
        # Create chunks
        chunks = DocumentProcessor.create_chunks(text, document_id, metadata or {})
        
        # Generate embeddings
        chunks = DocumentProcessor.generate_embeddings(chunks)
        
        # Store in vector database (implement this based on your vector DB)
        # store_in_vector_db(chunks)
        
        # Store in regular database (implement this based on your DB)
        # store_in_database(chunks)
        
        logger.info(f"Successfully processed document {document_id} into {len(chunks)} chunks")
        
        return {
            'document_id': document_id,
            'status': 'completed',
            'chunks_count': len(chunks),
            'processing_time': 0.0  # Add timing if needed
        }
        
    except Exception as e:
        logger.error(f"Error processing document {document_id}: {e}")
        self.retry(countdown=60, max_retries=3)
        return {
            'document_id': document_id,
            'status': 'failed',
            'chunks_count': 0,
            'processing_time': 0.0,
            'error': str(e)
        }

# API endpoints
@app.post("/process", response_model=ProcessingResult)
async def process_document(
    request: DocumentProcessingRequest,
    background_tasks: BackgroundTasks
):
    """Process a document"""
    try:
        # Validate file exists
        if not os.path.exists(request.file_path):
            raise HTTPException(status_code=404, detail="File not found")
        
        # Queue processing task
        task = process_document_task.delay(
            request.document_id,
            request.file_path,
            request.file_type,
            request.metadata
        )
        
        return ProcessingResult(
            document_id=request.document_id,
            status="queued",
            chunks_count=0,
            processing_time=0.0
        )
        
    except Exception as e:
        logger.error(f"Error queuing document processing: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/status/{document_id}")
async def get_processing_status(document_id: str):
    """Get processing status for a document"""
    # Implement status checking logic
    return {"document_id": document_id, "status": "processing"}

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "embedding_model": EMBEDDING_MODEL,
        "chunk_size": CHUNK_SIZE,
        "supported_formats": DocumentProcessor.SUPPORTED_FORMATS
    }

if __name__ == "__main__":
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8001,
        reload=True
    )
