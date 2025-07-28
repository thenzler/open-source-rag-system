"""
Document Processing Service
Handles business logic for document management operations
"""
import logging
import os
import hashlib
import mimetypes
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import asyncio

from ..repositories.interfaces import IDocumentRepository, IVectorSearchRepository
from ..repositories.audit_repository import SwissAuditRepository
from ..models.api_models import DocumentResponse, DocumentUpdate
try:
    from ..config.config import config
except ImportError:
    try:
        from config.config import config
    except ImportError:
        config = None

logger = logging.getLogger(__name__)

class DocumentProcessingService:
    """Service for document processing business logic"""
    
    def __init__(
        self,
        doc_repo: IDocumentRepository,
        vector_repo: IVectorSearchRepository,
        audit_repo: SwissAuditRepository
    ):
        self.doc_repo = doc_repo
        self.vector_repo = vector_repo
        self.audit_repo = audit_repo
        
        # File validation settings
        self.max_file_size = getattr(config, 'MAX_FILE_SIZE', 50 * 1024 * 1024)  # 50MB
        self.allowed_extensions = {'.pdf', '.docx', '.txt', '.md', '.csv', '.xlsx'}
        self.allowed_mime_types = {
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'text/plain',
            'text/markdown',
            'text/csv',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        }
        
        # Content filtering settings
        self.problematic_keywords = [
            'zero-hallucination', 'guidelines for following', 'only use information',
            'additional guidelines', 'training instructions', 'quelels'
        ]
        self.bio_waste_keywords = [
            'bioabfall', 'bio waste', 'organic waste', 'kompost', 'grünabfall',
            'küchenabfälle', 'obst', 'gemüse', 'fruit', 'vegetable', 'food waste'
        ]
    
    async def validate_upload(self, filename: str, content: bytes, content_type: str) -> Tuple[bool, str]:
        """Validate uploaded file"""
        try:
            # Check filename
            if not filename or filename.strip() == "":
                return False, "Filename cannot be empty"
            
            # Check file extension
            file_ext = Path(filename).suffix.lower()
            if file_ext not in self.allowed_extensions:
                return False, f"File type {file_ext} not allowed. Allowed: {', '.join(self.allowed_extensions)}"
            
            # Check file size
            if len(content) > self.max_file_size:
                return False, f"File size {len(content)} exceeds maximum {self.max_file_size} bytes"
            
            # Check MIME type
            detected_type = mimetypes.guess_type(filename)[0]
            if detected_type not in self.allowed_mime_types:
                return False, f"MIME type {detected_type} not allowed"
            
            # Basic content validation
            if len(content) == 0:
                return False, "File cannot be empty"
            
            # Security: Check for suspicious filenames
            if any(char in filename for char in ['..', '/', '\\', '<', '>', ':', '"', '|', '?', '*']):
                return False, "Filename contains invalid characters"
            
            return True, "Valid file"
            
        except Exception as e:
            logger.error(f"File validation error: {e}")
            return False, f"Validation failed: {str(e)}"
    
    def _generate_file_hash(self, content: bytes) -> str:
        """Generate SHA-256 hash of file content"""
        return hashlib.sha256(content).hexdigest()
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename for safe storage"""
        # Remove path components
        filename = os.path.basename(filename)
        
        # Replace problematic characters
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        
        # Limit length
        if len(filename) > 255:
            name, ext = os.path.splitext(filename)
            filename = name[:250-len(ext)] + ext
        
        return filename
    
    async def process_upload(
        self, 
        filename: str, 
        content: bytes, 
        content_type: str,
        uploader_id: Optional[str] = None
    ) -> DocumentResponse:
        """Process document upload with full business logic"""
        try:
            # Validate upload
            is_valid, validation_message = await self.validate_upload(filename, content, content_type)
            if not is_valid:
                # Note: Audit logging would go here in production
                # await self.audit_repo.log_event(audit_entry)
                raise ValueError(validation_message)
            
            # Sanitize filename
            safe_filename = self._sanitize_filename(filename)
            
            # Generate file hash for deduplication
            file_hash = self._generate_file_hash(content)
            
            # Check for duplicates
            existing = await self.doc_repo.find_by_hash(file_hash)
            if existing:
                logger.info(f"Duplicate file detected: {safe_filename}")
                # Note: Audit logging would go here in production
                return DocumentResponse(
                    id=existing.id,
                    filename=existing.filename,
                    size=existing.file_size,
                    content_type=existing.content_type,
                    status="duplicate",
                    message="File already exists"
                )
            
            # Save file to storage
            upload_dir = Path(config.UPLOAD_DIR if config and hasattr(config, 'UPLOAD_DIR') else 'storage/uploads')
            upload_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate unique filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            stored_filename = f"{timestamp}_{safe_filename}"
            file_path = upload_dir / stored_filename
            
            # Write file
            with open(file_path, 'wb') as f:
                f.write(content)
            
            # Create document record
            from ..repositories.models import Document, DocumentStatus
            
            document = Document(
                filename=safe_filename,
                original_filename=filename,
                file_path=str(file_path),
                file_size=len(content),
                content_type=content_type,
                uploader=uploader_id or 'anonymous',
                upload_timestamp=datetime.now(),
                status=DocumentStatus.UPLOADING,
                metadata={'file_hash': file_hash}
            )
            
            # Store in repository
            document = await self.doc_repo.create(document)
            
            # Start async processing
            asyncio.create_task(self._process_document_async(document.id, file_path))
            
            # Note: Audit logging would go here in production
            logger.info(f"Document uploaded successfully: {safe_filename} (ID: {document.id})")
            
            return DocumentResponse(
                id=document.id,
                filename=document.filename,
                size=document.file_size,
                content_type=document.content_type,
                status="uploaded",
                message="Document uploaded successfully"
            )
            
        except Exception as e:
            logger.error(f"Document upload processing failed: {e}")
            # Note: Audit logging would go here in production
            raise
    
    async def _process_document_async(self, document_id: int, file_path: Path):
        """Process document in background (chunking, embedding, etc.)"""
        try:
            # Update status to processing
            await self.doc_repo.update_status(document_id, 'processing')
            
            # Extract text from document
            text_content = await self._extract_text(file_path)
            if not text_content:
                raise ValueError("No text content extracted from document")
            
            # Update document with text content
            await self.doc_repo.update(document_id, {'text_content': text_content})
            
            # Create chunks
            chunks = self._create_chunks(text_content, chunk_size=500, overlap=50)
            logger.info(f"Created {len(chunks)} chunks for document {document_id}")
            
            # Store chunks in database
            chunk_ids = await self._store_chunks(document_id, chunks)
            
            # Generate embeddings for chunks
            embeddings = await self._generate_embeddings(chunks)
            
            # Store embeddings in database
            embedding_records = await self._store_embeddings(document_id, chunk_ids, embeddings)
            
            # Update vector index
            await self.vector_repo.add_to_index(embedding_records)
            
            # Update document status and counts
            await self.doc_repo.update(document_id, {
                'status': 'completed',
                'chunk_count': len(chunks),
                'embedding_count': len(embeddings),
                'completion_timestamp': datetime.now()
            })
            
            logger.info(f"Document {document_id} processed successfully with {len(chunks)} chunks")
            
        except Exception as e:
            logger.error(f"Document processing failed for {document_id}: {e}")
            await self.doc_repo.update_status(document_id, 'failed')
    
    async def get_document_details(self, document_id: int) -> Dict[str, Any]:
        """Get detailed document information"""
        try:
            document = await self.doc_repo.get_by_id(document_id)
            if not document:
                raise ValueError(f"Document {document_id} not found")
            
            return {
                "document": {
                    "id": document.id,
                    "filename": document.filename,
                    "original_filename": getattr(document, 'original_filename', document.filename),
                    "size": document.file_size,
                    "content_type": document.content_type,
                    "status": document.status,
                    "upload_date": document.upload_date.isoformat() if document.upload_date else None,
                    "uploader": getattr(document, 'uploader', 'unknown'),
                    "hash": getattr(document, 'file_hash', None)
                }
            }
            
        except Exception as e:
            logger.error(f"Error getting document details: {e}")
            raise
    
    async def update_document(self, document_id: int, updates: DocumentUpdate) -> Dict[str, Any]:
        """Update document metadata"""
        try:
            # Get existing document
            document = await self.doc_repo.get_by_id(document_id)
            if not document:
                raise ValueError(f"Document {document_id} not found")
            
            # Prepare update data
            update_data = updates.dict(exclude_unset=True)
            
            # Apply updates
            updated_document = await self.doc_repo.update(document_id, update_data)
            
            # Note: Audit logging would go here in production
            logger.info(f"Document {document_id} updated successfully")
            
            return {
                "message": "Document updated successfully",
                "document": {
                    "id": updated_document.id,
                    "updates": update_data
                }
            }
            
        except Exception as e:
            logger.error(f"Error updating document: {e}")
            raise
    
    async def delete_document(self, document_id: int) -> Dict[str, Any]:
        """Delete document and cleanup"""
        try:
            # Get document for cleanup
            document = await self.doc_repo.get_by_id(document_id)
            if not document:
                raise ValueError(f"Document {document_id} not found")
            
            # Delete from vector store if exists
            try:
                await self.vector_repo.delete_document(document_id)
            except Exception as e:
                logger.warning(f"Vector cleanup failed for document {document_id}: {e}")
            
            # Delete physical file
            if hasattr(document, 'file_path') and document.file_path:
                try:
                    file_path = Path(document.file_path)
                    if file_path.exists():
                        file_path.unlink()
                        logger.info(f"Deleted file: {file_path}")
                except Exception as e:
                    logger.warning(f"File cleanup failed: {e}")
            
            # Delete from repository
            await self.doc_repo.delete(document_id)
            
            # Note: Audit logging would go here in production
            logger.info(f"Document {document_id} deleted successfully")
            
            return {
                "message": f"Document {document_id} deleted successfully"
            }
            
        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            raise
    
    async def get_download_path(self, document_id: int) -> Path:
        """Get file path for document download"""
        try:
            document = await self.doc_repo.get_by_id(document_id)
            if not document:
                raise ValueError(f"Document {document_id} not found")
            
            if not hasattr(document, 'file_path') or not document.file_path:
                raise ValueError(f"No file path for document {document_id}")
            
            file_path = Path(document.file_path)
            if not file_path.exists():
                raise ValueError(f"File not found: {file_path}")
            
            return file_path
            
        except Exception as e:
            logger.error(f"Error getting download path: {e}")
            raise
    
    async def _extract_text(self, file_path: Path) -> str:
        """Extract text from document based on file type"""
        try:
            file_ext = file_path.suffix.lower()
            
            if file_ext == '.txt' or file_ext == '.md':
                # Read plain text files
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            
            elif file_ext == '.pdf':
                # Extract from PDF
                try:
                    import PyPDF2
                    text = ""
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
                    return text.strip()
                except ImportError:
                    logger.warning("PyPDF2 not installed, trying fallback")
                    # Fallback: just return filename
                    return f"PDF Document: {file_path.name}"
            
            elif file_ext == '.docx':
                # Extract from Word documents
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(file_path)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    return text.strip()
                except ImportError:
                    logger.warning("python-docx not installed, trying fallback")
                    return f"Word Document: {file_path.name}"
            
            elif file_ext in ['.csv', '.xlsx']:
                # Extract from spreadsheets
                try:
                    import pandas as pd
                    if file_ext == '.csv':
                        df = pd.read_csv(file_path)
                    else:
                        df = pd.read_excel(file_path)
                    # Convert dataframe to text
                    return df.to_string()
                except ImportError:
                    logger.warning("pandas not installed, trying fallback")
                    return f"Spreadsheet: {file_path.name}"
            
            else:
                # Unknown file type - try to read as text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
                    
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            # Return at least the filename so we have something
            return f"Document: {file_path.name}"
    
    def _create_chunks(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Split text into overlapping chunks"""
        if not text:
            return []
        
        chunks = []
        words = text.split()
        
        if len(words) <= chunk_size:
            # Text is small enough to be a single chunk
            return [text]
        
        # Create overlapping chunks
        for i in range(0, len(words), chunk_size - overlap):
            chunk_words = words[i:i + chunk_size]
            chunk_text = ' '.join(chunk_words)
            if chunk_text.strip():  # Only add non-empty chunks
                chunks.append(chunk_text)
        
        return chunks
    
    async def _store_chunks(self, document_id: int, chunks: List[str]) -> List[int]:
        """Store text chunks in database"""
        chunk_ids = []
        
        # Get chunk repository (we need to add this to interfaces)
        # For now, store in SQLite directly
        import sqlite3
        db_path = getattr(config, 'DATABASE_PATH', None) if config else None
        conn = sqlite3.connect(db_path or 'data/rag_database.db')
        
        try:
            for i, chunk_text in enumerate(chunks):
                cursor = conn.execute("""
                    INSERT INTO chunks (document_id, chunk_index, text, 
                                      character_count, word_count)
                    VALUES (?, ?, ?, ?, ?)
                """, (
                    document_id,
                    i,
                    chunk_text,
                    len(chunk_text),
                    len(chunk_text.split())
                ))
                chunk_ids.append(cursor.lastrowid)
            
            conn.commit()
            return chunk_ids
            
        finally:
            conn.close()
    
    async def _generate_embeddings(self, chunks: List[str]) -> List[List[float]]:
        """Generate embeddings for text chunks"""
        try:
            # Use sentence transformers
            from sentence_transformers import SentenceTransformer
            
            # Initialize model (this should be cached in production)
            model = SentenceTransformer('all-MiniLM-L6-v2')
            
            # Generate embeddings
            embeddings = model.encode(chunks, convert_to_numpy=True)
            
            return embeddings.tolist()
            
        except ImportError:
            logger.error("sentence-transformers not installed")
            # Return dummy embeddings
            import random
            return [[random.random() for _ in range(384)] for _ in chunks]
    
    async def _store_embeddings(self, document_id: int, chunk_ids: List[int], 
                               embeddings: List[List[float]]) -> List:
        """Store embeddings in database"""
        from ..repositories.models import Embedding
        
        embedding_records = []
        
        # Store in SQLite directly for now
        import sqlite3
        import pickle
        import gzip
        
        db_path = getattr(config, 'DATABASE_PATH', None) if config else None
        conn = sqlite3.connect(db_path or 'data/rag_database.db')
        
        try:
            for chunk_id, embedding_vector in zip(chunk_ids, embeddings):
                # Compress embedding vector
                compressed = gzip.compress(pickle.dumps(embedding_vector))
                
                cursor = conn.execute("""
                    INSERT INTO embeddings (chunk_id, embedding_data, 
                                          embedding_model, dimensions)
                    VALUES (?, ?, ?, ?)
                """, (
                    chunk_id,
                    compressed,
                    'all-MiniLM-L6-v2',
                    len(embedding_vector)
                ))
                
                # Create Embedding object for vector index
                embedding_obj = Embedding(
                    id=cursor.lastrowid,
                    chunk_id=chunk_id,
                    document_id=document_id,
                    embedding_vector=embedding_vector,
                    embedding_model='all-MiniLM-L6-v2',
                    vector_dimension=len(embedding_vector)
                )
                embedding_records.append(embedding_obj)
            
            conn.commit()
            return embedding_records
            
        finally:
            conn.close()
    
    def analyze_document_content(self, text_content: str) -> Dict[str, Any]:
        """Analyze document content for problematic patterns"""
        content_lower = text_content.lower()
        
        # Check for problematic content
        problematic_score = sum(1 for keyword in self.problematic_keywords if keyword in content_lower)
        
        # Check for bio waste content
        bio_waste_score = sum(1 for keyword in self.bio_waste_keywords if keyword in content_lower)
        
        # Check for encoding issues
        corruption_score = text_content.count('�')
        
        # Classify content type
        content_type = "unknown"
        if bio_waste_score >= 2:
            content_type = "bio_waste"
        elif problematic_score > 0:
            content_type = "training_instructions"
        elif any(cs in content_lower for cs in ['javascript', 'programming', 'software']):
            content_type = "computer_science"
        
        is_problematic = (
            problematic_score > 0 or 
            (corruption_score > 10) or 
            (len(text_content.strip()) < 100 and content_type != "bio_waste")
        )
        
        return {
            "content_type": content_type,
            "is_problematic": is_problematic,
            "bio_waste_score": bio_waste_score,
            "problematic_score": problematic_score,
            "corruption_score": corruption_score,
            "content_length": len(text_content),
            "recommendation": "reject" if is_problematic else "accept"
        }
    
    async def validate_document_content(self, text_content: str) -> Tuple[bool, str, Dict[str, Any]]:
        """Validate document content and return analysis"""
        analysis = self.analyze_document_content(text_content)
        
        if analysis["is_problematic"]:
            reasons = []
            if analysis["problematic_score"] > 0:
                reasons.append("contains training instructions")
            if analysis["corruption_score"] > 10:
                reasons.append("has encoding corruption")
            if analysis["content_length"] < 100:
                reasons.append("content too short")
            
            return False, f"Document rejected: {', '.join(reasons)}", analysis
        
        if analysis["content_type"] == "bio_waste":
            return True, "Bio waste document accepted", analysis
        elif analysis["content_type"] == "unknown" and len(text_content.strip()) > 200:
            return True, "General document accepted", analysis
        else:
            return False, "Document type not suitable for bio waste RAG system", analysis